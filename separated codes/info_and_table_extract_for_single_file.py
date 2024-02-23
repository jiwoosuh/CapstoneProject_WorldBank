import os
import re
import csv
from docx import Document

def extract_info_from_docx(docx_file):
    doc = Document(docx_file)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    respondent_id_match = re.search(r'Respondent ID:\s*([^\n]+)', text)
    date_match = re.search(r'Date:\s*(\d{1,2}(?:st|nd|rd|th)?\s*\w+\s*\d{4}|[\d./]+)', text)
    week_match = re.search(r'WK\s+(\d+)', text, re.IGNORECASE)
    wag_match = re.search(r'(?:BL|ML)\d+-(\w+)-\w+', text)

    # Extract information if there are matches or return None
    respondent_id = respondent_id_match.group(1) if respondent_id_match else None
    date = date_match.group(1) if date_match else None
    week = week_match.group(1) if week_match else None
    wag = wag_match.group(1) if wag_match else None

    return respondent_id, date, week, wag

def convert_table_to_csv_file(input_docx, output_csv):
    # Read the Word document
    doc = Document(input_docx)

    # Define the header for CSV file
    csv_file_header = ['Respondent ID', 'Date', 'Week', 'Member Status', 'Transaction Nature', 'Transaction Type', 'Transaction Name', 'Transaction Amount', 'Transaction Comment']

    # Initialize the data for CSV file
    csv_file_data = []

    # Extract information from docx
    respondent_id, date, week, wag = extract_info_from_docx(input_docx)

    transaction_type = "Variable"
    for table in doc.tables:
        rows = table.rows
        n = len(rows)

        # Income section
        for i in range(n):
            transaction_nature = "Income"
            transaction_name = rows[i].cells[0].text.strip()  # The first column is transaction name
            transaction_amount = rows[i].cells[1].text.strip()  # The second column is transaction amount
            if transaction_name == "Fixed weekly income":
                transaction_type = "Fixed"
            elif transaction_name == "Variable weekly income":
                transaction_type = "Variable"

            # Skip specific content
            if transaction_name in ["Fixed weekly income", "Variable weekly income", "Total:", "Total", "Comments", ""]:
                continue

            # Combine the extracted data into a row and add it to the CSV file data
            csv_file_data.append([respondent_id, date, week, wag, transaction_nature, transaction_type, transaction_name, transaction_amount, ""])

        # Expenditure section
        for i in range(n):
            transaction_nature = "Expenditure"
            transaction_name = rows[i].cells[2].text.strip()  # The third column is transaction name
            transaction_amount = rows[i].cells[3].text.strip()  # The fourth column is transaction amount
            try:
                transaction_comment = rows[i].cells[4].text.strip()  # The fifth column is transaction comment
            except IndexError:
                transaction_comment = ""
            if transaction_name == "Fixed weekly income":
                transaction_type = "Fixed"
            elif transaction_name == "Variable weekly income":
                transaction_type = "Variable"

            # Skip specific content
            if transaction_name in ["Fixed weekly expenditure", "Variable weekly expenditure", "Total:", "Total", "Comments", ""]:
                continue

            # Combine the extracted data into a row and add it to the CSV file data
            csv_file_data.append([respondent_id, date, week, wag, transaction_nature, transaction_type, transaction_name, transaction_amount, transaction_comment])

    # Write the data to a CSV file
    with open(output_csv, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(csv_file_header)  # Write the header
        writer.writerows(csv_file_data)  # Write the data

# Specify the input Word document and output CSV file path
input_docx = 'Abia_Obingwa_Abayi Ward I_IDI_FD_WAG Member_Week 4.docx'
output_csv = 'combined.csv'

# Convert Table to CSV file
convert_table_to_csv_file(input_docx, output_csv)

print(f'Combined data saved to {output_csv}')
