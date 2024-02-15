from docx import Document
import pandas as pd
import os
import csv
#%%

def convert_docx_to_txt(docx_filename, txt_filename, output_dir=None):
    doc = Document(docx_filename)
    if output_dir:
        os.makedirs(output_dir,exist_ok=True)
        txt_filename = os.path.join(output_dir,txt_filename)

    with open(txt_filename, 'w', encoding='utf-8') as txt_file:
        for paragraph in doc.paragraphs:
            txt_file.write(paragraph.text + '\n')

def extract_tables_to_text(docx_filename, txt_filename, output_dir):
    doc = Document(docx_filename)

    if output_dir:
        os.makedirs(output_dir,exist_ok=True)
        txt_filename = os.path.join(output_dir,txt_filename)

    with open(txt_filename, 'w', encoding='utf-8', newline='') as txt_file:
        csv_writer = csv.writer(txt_file, delimiter='\t')

        for table_index, table in enumerate(doc.tables, start=1):
            txt_file.write(f"Table {table_index}:\n")

            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                csv_writer.writerow(row_data)

            txt_file.write('\n')

def extract_tables_to_dataframe(docx_filename, output_dir=None):
    doc = Document(docx_filename)
    table_dataframes = []

    for table_index, table in enumerate(doc.tables, start=1):
        table_data = []

        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)

        df = pd.DataFrame(table_data, columns=[f'Column_{i+1}' for i in range(len(table_data[0]))])
        table_dataframes.append(df)

    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

        for idx, df in enumerate(table_dataframes, start=1):
            output_filename = os.path.join(output_dir, f'table_{idx}.csv')
            df.to_csv(output_filename, index=False)

    return table_dataframes


docx_file_path = 'Financial_Diaries/Financial Diaries/Midline Financial Diaries/Abia/ISIALA NGWA/NON WAG/Abia_Isiala Ngwa North_Amasaa Ward _Non-WAG Member_Week 1.docx'
output_text_file = 'output.txt'
output_directory = 'Code/output_text'
# convert_docx_to_txt(docx_file_path, output_text_file, output_dir=output_directory)
# extract_tables_to_text(docx_file_path, output_text_file, output_dir=output_directory)
pd_table = extract_tables_to_dataframe(docx_file_path, output_dir=output_directory)
print(pd_table)