import os
import re
import csv
from docx import Document

# Function to get all file locations in a directory and its subdirectories
def get_file_locations(folder):
    file_locations = []
    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.endswith('.docx'):  # Check if file ends with ".docx"
                file_locations.append(os.path.join(root, file))
    return file_locations
def extract_info_from_docx(docx_file):
    if docx_file == "FINANCIAL DIARY FOR NON WAG MEMBER BALI LGA (URBAN) 15122021.docx":
        return None  # Skip processing this file
    doc = Document(docx_file)
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    respondent_id_match = re.search(r'(ML\d+.*WK\s*\d+|BL\d+.*WK\s*\d+)', text)
    # date_match = re.search(r'(?i)(?:Date|DATE):\s*((?:\d{1,2}(?:st|nd|rd|th)?\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})|(?:\d{1,2}/\d{1,2}/\d{2,4})|(?:\d{1,2}[./]\d{1,2}[./]\d{2,4}))\s*[_]*', text)
    date_match = re.search(r'(?i)(?:Date|DATE):\s*((?:\d{1,2}(?:st|nd|rd|th)?\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})|(?:\d{1,2}/\d{1,2}/\d{2,4})|(?:\d{1,2}[./]\d{1,2}[./]\d{2,4})(?:\s*[_]*)*)',
        text)

    week_match = re.search(r'-WK\s*(\d+)', text, re.IGNORECASE)

    # Extract information if there are matches or return None
    respondent_id = respondent_id_match.group(1) if respondent_id_match else None
    date = date_match.group(1) if date_match else None
    week = week_match.group(1) if week_match else None

# Extracting other information from file path
    parts = docx_file.split(os.path.sep)
    folders = parts[-6:]
    if 'FDs' in folders:
        folders.remove('FDs')
    if 'Financial Diaries_3' in folders:
        folders.remove('Financial Diaries_3')
    if 'Data' in folders:
        folders.remove('Data')

    # Extracting other information from file path
    fd_name = folders[-5]
    state = folders[-4]
    region = folders[-3]
    member_state = folders[-2]
    file_name = folders[-1]

    return fd_name, state, region, member_state, file_name, respondent_id, date, week

def convert_table_to_csv_file(input_docx, csv_file_header):
    # Read the Word document
    doc = Document(input_docx)

    # Initialize the data for CSV file
    csv_file_data = []

    # Extract information from docx
    fd_name, state, region, member_state, file_name, respondent_id, date, week = extract_info_from_docx(input_docx)

    transaction_type = "Variable"
    for table in doc.tables:
        rows = table.rows
        n = len(rows)

        # Income section
        for i in range(n):
            transaction_nature = "Income"
            transaction_name = rows[i].cells[0].text.strip()  # The first column is transaction name
            transaction_amount = rows[i].cells[1].text.strip()  # The second column is transaction amount
            if transaction_name == "Fixed weekly income" or transaction_name == "FIXED WEEKLY INCOME":
                transaction_type = "Fixed"
            elif transaction_name == "Variable weekly income" or transaction_name == "VARIABLE INCOME":
                transaction_type = "Variable"

            # Skip specific content
            if transaction_name in ["Fixed weekly income", "Variable weekly income", "Total:", "Total", "Comments", "", "#"]:
                continue

            # Combine the extracted data into a row and add it to the CSV file data
            csv_file_data.append(
                [fd_name, state, region, member_state, file_name, respondent_id, date, week,
                 transaction_type,
                 transaction_nature, transaction_name, transaction_amount])

        # Expenditure section
        for i in range(n):
            transaction_nature = "Expenditure"
            transaction_name = rows[i].cells[2].text.strip()  # The third column is transaction name
            transaction_amount = rows[i].cells[3].text.strip()  # The fourth column is transaction amount
            try:
                transaction_comment = rows[i].cells[4].text.strip()  # The fifth column is transaction comment
            except IndexError:
                transaction_comment = ""
            if transaction_name == "Fixed weekly expenditure" or transaction_name == "FIXED WEEKLY EXPENDITURE":
                transaction_type = "Fixed"
            elif transaction_name == "Variable weekly expenditure" or transaction_name == "VARIABLE WEEKLY EXPENDITURE":
                transaction_type = "Variable"

            # Skip specific content
            if transaction_name in ["Fixed weekly expenditure", "Variable weekly expenditure", "Total:", "Total", "Comments", "", "#"]:
                continue

            # Combine the extracted data into a row and add it to the CSV file data
            csv_file_data.append(
                [fd_name, state, region, member_state, file_name, respondent_id, date, week,
                 transaction_type,
                 transaction_nature, transaction_name, transaction_amount, transaction_comment])

    return csv_file_data

# Folder to search for files
folder = 'Data'

# Get file locations
file_locations = get_file_locations(folder)

# Initialize CSV file header
csv_file_header = ['FD_Name', 'State', 'Region', 'Member_Status', 'File_Name', 'Respondent ID', 'Date', 'Week', 'Transaction_Nature', 'Transaction_Type',  'Transaction_Name', 'Transaction_Amount', 'Transaction_Comment']

# Initialize the data for CSV file
combined_csv_data = []

# Convert each table to CSV data
for docx_file in file_locations:
    print(f'Processing file: {docx_file}')
    csv_data = convert_table_to_csv_file(docx_file, csv_file_header)
    combined_csv_data.extend(csv_data)

# Write the combined data to a CSV file
combined_output_csv = 'combined_output.csv'
with open(combined_output_csv, 'w', newline='', encoding='utf-8') as csvfile:
    writer = csv.writer(csvfile)
    writer.writerow(csv_file_header)  # Write the header
    writer.writerows(combined_csv_data)  # Write the data

print(f'Combined data saved to {combined_output_csv}')
